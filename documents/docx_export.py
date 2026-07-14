"""
Экспорт документов в Word (.docx) — PRO-фича (раздел Модуль A ТЗ: «Скачивание в
формате PDF + исходник Word (опционально для PRO-тарифа)»).

Содержание зеркалит PDF-шаблоны (templates/documents/pdf/), но строится напрямую
через python-docx, а не конвертацией HTML — так надёжнее и предсказуемее по вёрстке.
"""

from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm, RGBColor

NAVY = RGBColor(0x1E, 0x3A, 0x5F)

_MESYATSY_RU = {
    1: 'января', 2: 'февраля', 3: 'марта', 4: 'апреля', 5: 'мая', 6: 'июня',
    7: 'июля', 8: 'августа', 9: 'сентября', 10: 'октября', 11: 'ноября', 12: 'декабря',
}


def _data_ru(dt) -> str:
    """Форматирует дату по-русски без зависимости от системной локали."""
    return f'{dt.day} {_MESYATSY_RU[dt.month]} {dt.year}'


def _setup_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    return doc


def _title(doc, text, subtitle=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = NAVY
    if subtitle:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run(subtitle)
        run2.italic = True
        run2.font.size = Pt(10)
    doc.add_paragraph()


def _section_title(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = NAVY
    p.space_before = Pt(10)


def _brigada_header(doc, brigada):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(brigada.nazvanie)
    run.bold = True
    run.font.color.rgb = NAVY
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p2.add_run(f'Тел: {brigada.telefon}').font.size = Pt(9)
    if brigada.rekvizity:
        p3 = doc.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p3.add_run(brigada.rekvizity).font.size = Pt(9)
    doc.add_paragraph()


def _signature_block(doc, brigada_name, zakazchik_name, label_left='Исполнитель:', label_right='Заказчик:'):
    doc.add_paragraph()
    table = doc.add_table(rows=3, cols=2)
    table.autofit = True
    table.cell(0, 0).paragraphs[0].add_run(label_left).bold = True
    table.cell(0, 1).paragraphs[0].add_run(label_right).bold = True
    table.cell(1, 0).paragraphs[0].add_run(brigada_name)
    table.cell(1, 1).paragraphs[0].add_run(zakazchik_name)
    table.cell(2, 0).paragraphs[0].add_run('_______________ / подпись')
    table.cell(2, 1).paragraphs[0].add_run('_______________ / подпись')


def _raboty_table(doc, headers, rows, total_label, total_value):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    total_row = table.add_row().cells
    total_row[0].text = total_label
    total_row[0].paragraphs[0].runs[0].bold = True
    total_row[-1].text = str(total_value)
    total_row[-1].paragraphs[0].runs[0].bold = True
    doc.add_paragraph()


def _dogovor_docx(d, brigada) -> Document:
    doc = _setup_document()
    _brigada_header(doc, brigada)
    _title(doc, f'ДОГОВОР ПОДРЯДА №{d.nomer}', 'на выполнение ремонтно-строительных работ')

    doc.add_paragraph(f'г. {brigada.region or "—"}          {_data_ru(d.data_sozdaniya)} г.')
    doc.add_paragraph(
        f'{brigada.nazvanie}, именуемый в дальнейшем «Исполнитель», с одной стороны, и '
        f'{d.zakazchik} (тел. {d.zakazchik_telefon}), именуемый в дальнейшем «Заказчик», с другой '
        f'стороны, вместе именуемые «Стороны», заключили настоящий Договор о нижеследующем:'
    )

    _section_title(doc, '1. Предмет договора')
    doc.add_paragraph(
        f'Исполнитель обязуется по заданию Заказчика выполнить ремонтно-строительные работы по '
        f'адресу: {d.adres_obekta} (далее — «Объект»), а Заказчик обязуется принять результат работ '
        f'и оплатить его в порядке и в сроки, установленные настоящим Договором.'
    )

    _section_title(doc, '2. Цена договора и порядок оплаты')
    doc.add_paragraph(f'Общая стоимость работ по настоящему Договору составляет {d.summa} ₽. Оплата производится поэтапно:')
    _raboty_table(
        doc,
        ['Этап оплаты', 'Доля', 'Сумма, ₽', 'Условие'],
        [
            ['Аванс', '30%', str(d.avans_platezh), 'При подписании Договора'],
            ['Промежуточный платёж', '40%', str(d.promezhutochny_platezh), 'После завершения основного объёма работ'],
            ['Окончательный расчёт', '30%', str(d.okonchatelny_platezh), 'По подписании Акта выполненных работ'],
        ],
        'ИТОГО', str(d.summa),
    )

    _section_title(doc, '3. Сроки выполнения работ')
    nachalo = d.srok_nachala.strftime('%d.%m.%Y') if d.srok_nachala else 'согласовывается дополнительно'
    okonchanie = d.srok_okonchania.strftime('%d.%m.%Y') if d.srok_okonchania else 'согласовывается дополнительно'
    doc.add_paragraph(f'Начало работ: {nachalo}. Плановое окончание работ: {okonchanie}.')

    _section_title(doc, '4. Ответственность сторон')
    doc.add_paragraph(
        '4.1. В случае нарушения Заказчиком сроков оплаты Заказчик уплачивает Исполнителю пеню '
        'в размере 0,1% от неоплаченной суммы за каждый день просрочки, но не более 10% от суммы платежа.'
    )
    doc.add_paragraph(
        '4.2. В случае нарушения Исполнителем согласованных сроков выполнения работ по его вине, '
        'Исполнитель уплачивает Заказчику пеню в размере 0,1% от стоимости этапа за каждый день просрочки, '
        'но не более 10% от стоимости такого этапа.'
    )

    _section_title(doc, '5. Порядок изменения объёмов работ')
    doc.add_paragraph(
        'Любое изменение объёма, состава или стоимости работ оформляется дополнительным соглашением, '
        'подписанным обеими Сторонами, до начала выполнения соответствующих дополнительных работ.'
    )

    _section_title(doc, '6. Гарантийные обязательства')
    doc.add_paragraph(
        'Исполнитель предоставляет гарантию качества выполненных работ сроком 1 (один) год с момента '
        'подписания итогового Акта выполненных работ.'
    )

    _section_title(doc, '7. Порядок расторжения договора')
    doc.add_paragraph(
        'Договор может быть расторгнут по соглашению Сторон, а также в одностороннем порядке при '
        'существенном нарушении его условий другой Стороной, с письменным уведомлением не менее чем за 5 дней.'
    )

    _signature_block(doc, brigada.nazvanie, d.zakazchik)
    return doc


def _akt_vkr_docx(d, brigada) -> Document:
    doc = _setup_document()
    _brigada_header(doc, brigada)
    _title(doc, f'АКТ №{d.nomer}', 'о выполненных работах (форма ВКР)')

    data_akta = d.srok_okonchania.strftime('%d.%m.%Y') if d.srok_okonchania else d.data_sozdaniya.strftime('%d.%m.%Y')
    doc.add_paragraph(f'Объект: {d.adres_obekta}     Дата: {data_akta}')
    doc.add_paragraph(
        f'Исполнитель {brigada.nazvanie} и Заказчик {d.zakazchik} составили настоящий Акт о том, что '
        f'выполнены и приняты следующие работы:'
    )

    rows = [[i + 1, p.nazvanie, p.edinica, p.kolvo, p.cena, p.summa] for i, p in enumerate(d.pozicii.all())]
    _raboty_table(
        doc, ['№', 'Наименование', 'Ед.', 'Кол-во', 'Цена, ₽', 'Сумма, ₽'],
        rows, 'ИТОГО', str(d.itogo_po_poziciyam),
    )
    doc.add_paragraph('Работы выполнены в полном объёме, в согласованные сроки, с надлежащим качеством.')
    _signature_block(doc, brigada.nazvanie, d.zakazchik)
    return doc


def _raspiska_docx(d, brigada) -> Document:
    doc = _setup_document()
    _brigada_header(doc, brigada)
    _title(doc, 'РАСПИСКА', 'о получении аванса по договору подряда')

    doc.add_paragraph(f'г. {brigada.region or "—"}          {_data_ru(d.data_sozdaniya)} г.')
    doc.add_paragraph(
        f'Я, представитель {brigada.nazvanie} (тел. {brigada.telefon}), подтверждаю получение от '
        f'{d.zakazchik} ({d.zakazchik_telefon}) денежных средств в размере {d.avans_summa} ₽ в качестве '
        f'аванса по договору подряда на выполнение работ по адресу: {d.adres_obekta}.'
    )

    _section_title(doc, 'Условие невозврата аванса')
    doc.add_paragraph(
        'В случае отказа Заказчика от исполнения договора по собственной инициативе, аванс возврату '
        'не подлежит, поскольку является компенсацией затрат Исполнителя на подготовку к выполнению работ.'
    )
    doc.add_paragraph(
        'В случае отказа Исполнителя от выполнения работ либо существенного нарушения условий договора, '
        'аванс подлежит возврату Заказчику в полном объёме в течение 10 рабочих дней.'
    )

    _signature_block(doc, brigada.nazvanie, d.zakazchik, 'Получил (Исполнитель):', 'Передал (Заказчик):')
    return doc


def _akt_priemki_docx(d, brigada) -> Document:
    doc = _setup_document()
    _brigada_header(doc, brigada)
    _title(doc, f'АКТ №{d.nomer}', 'приёмки этапа выполненных работ')

    doc.add_paragraph(f'Объект: {d.adres_obekta}')
    doc.add_paragraph(f'Этап работ: {d.etap_nazvanie}')
    doc.add_paragraph(f'Дата: {d.data_sozdaniya:%d.%m.%Y}')

    _section_title(doc, 'Результаты осмотра')
    for item in d.checklist:
        doc.add_paragraph(item, style='List Bullet')

    _section_title(doc, 'Защита от переделок')
    doc.add_paragraph(
        'Подписание настоящего Акта означает окончательную приёмку Заказчиком указанного этапа работ '
        'без права требования его безвозмездной переделки по основаниям, не зафиксированным письменно '
        'в настоящем Акте на момент его подписания.'
    )

    _signature_block(doc, brigada.nazvanie, d.zakazchik)
    return doc


_GENERATORS = {
    'dogovor': _dogovor_docx,
    'akt_vkr': _akt_vkr_docx,
    'raspiska': _raspiska_docx,
    'akt_priemki': _akt_priemki_docx,
}


def generate_docx(dokument) -> bytes:
    """Генерирует .docx для документа. Возвращает байты файла."""
    generator = _GENERATORS[dokument.tip]
    doc = generator(dokument, dokument.brigada)
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
